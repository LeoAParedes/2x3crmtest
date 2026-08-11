from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import re

src = Path(r'g:\Claude\2x3crmtest\docs\manual-usuario\manual-usuario-erp-supermercado.md')
out = Path(r'g:\Claude\2x3crmtest\docs\manual-usuario\Manual-Usuario-2x3-Operaciones-v1.1.0.docx')
text = src.read_text(encoding='utf-8')

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def clean_inline(s: str) -> str:
    s = s.replace('**', '')
    s = s.replace('`', '')
    s = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', s)
    return s.strip()


lines = text.splitlines()
i = 0
in_code = False
code_buf: list[str] = []
table_buf: list[str] = []


def flush_table() -> None:
    global table_buf
    if not table_buf:
        return
    rows: list[list[str]] = []
    for row in table_buf:
        if re.match(r'^\|?\s*-+', row):
            continue
        cells = [clean_inline(c) for c in row.strip().strip('|').split('|')]
        rows.append([c.strip() for c in cells])
    table_buf = []
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            val = row[c_idx] if c_idx < len(row) else ''
            cell = table.cell(r_idx, c_idx)
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True
    doc.add_paragraph()


while i < len(lines):
    line = lines[i]
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
        else:
            in_code = False
            p = doc.add_paragraph(clean_inline('\n'.join(code_buf)))
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            code_buf = []
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.strip().startswith('|'):
        table_buf.append(line)
        i += 1
        if i >= len(lines) or not lines[i].strip().startswith('|'):
            flush_table()
        continue

    flush_table()

    if not line.strip():
        i += 1
        continue

    if line.startswith('# '):
        doc.add_heading(clean_inline(line[2:]), level=0)
    elif line.startswith('## '):
        doc.add_heading(clean_inline(line[3:]), level=1)
    elif line.startswith('### '):
        doc.add_heading(clean_inline(line[4:]), level=2)
    elif line.startswith('#### '):
        doc.add_heading(clean_inline(line[5:]), level=3)
    elif line.startswith('> '):
        p = doc.add_paragraph(clean_inline(line[2:]))
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    elif re.match(r'^\d+\.\s', line.strip()):
        doc.add_paragraph(
            clean_inline(re.sub(r'^\d+\.\s*', '', line.strip())),
            style='List Number',
        )
    elif line.strip().startswith('- '):
        doc.add_paragraph(clean_inline(line.strip()[2:]), style='List Bullet')
    elif 'Texto alternativo' in line:
        p = doc.add_paragraph(clean_inline(line.strip().strip('*')))
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)
    elif line.strip() == '---':
        doc.add_paragraph('─' * 40)
    else:
        doc.add_paragraph(clean_inline(line))
    i += 1

flush_table()
doc.save(out)
print(f'Wrote {out} ({out.stat().st_size} bytes)')
